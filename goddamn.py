import os
import time
import random
from pymongo import MongoClient
import pandas as pd
from docx import Document
from dotenv import load_dotenv

# === Load env ===
load_dotenv(".env")
MONGO_URI = os.getenv("MONGO_URI")
DB_NAME = os.getenv("DB_NAME")

# === Connect to MongoDB ===
client = MongoClient(MONGO_URI)
db = client[DB_NAME]
summary_collection = db["trade_summary"]

# === Fetch data ===
query = {"user_id": "drl_bot"}
snapshots = list(summary_collection.find(query).sort("last_updated", -1).limit(50))

# === Prepare DataFrame ===
data = []
if snapshots:
    symbols = ["BTCUSDT", "ETHUSDT", "BNBUSDT"]

    for snap in snapshots:
        ts = snap.get("last_updated", 0)
        live_profit = snap.get("live_profit", 0)
        inventory = snap.get("current_inventory", {})

        # === Generate dummy trade details
        symbol = random.choice(symbols)
        if symbol == "BTCUSDT":
            opening_price = round(random.uniform(27000, 28000), 2)
        elif symbol == "ETHUSDT":
            opening_price = round(random.uniform(1800, 2200), 2)
        else:
            opening_price = round(random.uniform(250, 350), 2)
        closing_price = round(opening_price * random.uniform(0.995, 1.005), 2)
        stop_loss = round(opening_price * 0.98, 2)
        take_profit = round(opening_price * 1.02, 2)
        quantity = round(random.uniform(0.001, 0.005), 6)
        profit = round((closing_price - opening_price) * quantity, 2)

        data.append({
            "Time": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(ts)),
            "Symbol": symbol,
            "Side": random.choice(["BUY", "SELL"]),
            "Qty": quantity,
            "Open": opening_price,
            "Close": closing_price,
            "SL": stop_loss,
            "TP": take_profit,
            "Profit": profit
        })

df = pd.DataFrame(data)

# === Create Word document ===
doc = Document()
doc.add_heading("Trading Report", 0)
doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")

if not df.empty:
    total_trades = len(df)
    total_profit = df["Profit"].sum()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Total Trades: {total_trades}")
    doc.add_paragraph(f"Total Profit: ${total_profit:.2f}")

    doc.add_heading("Detailed Trades", level=2)
    table = doc.add_table(rows=1, cols=9)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time'
    hdr_cells[1].text = 'Symbol'
    hdr_cells[2].text = 'Side'
    hdr_cells[3].text = 'Qty'
    hdr_cells[4].text = 'Open'
    hdr_cells[5].text = 'Close'
    hdr_cells[6].text = 'SL'
    hdr_cells[7].text = 'TP'
    hdr_cells[8].text = 'Profit'

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Time"])
        cells[1].text = str(row["Symbol"])
        cells[2].text = str(row["Side"])
        cells[3].text = f"{row['Qty']:.6f}"
        cells[4].text = f"{row['Open']:.2f}"
        cells[5].text = f"{row['Close']:.2f}"
        cells[6].text = f"{row['SL']:.2f}"
        cells[7].text = f"{row['TP']:.2f}"
        cells[8].text = f"${row['Profit']:.2f}"

else:
    doc.add_paragraph("No portfolio snapshots found in trade_summary.")

# === Save with timestamp ===
filename = f"Detailed_Trading_Report_{time.strftime('%Y%m%d_%H%M%S')}.docx"
doc.save(filename)
print(f"✅ Report saved as {filename}")
